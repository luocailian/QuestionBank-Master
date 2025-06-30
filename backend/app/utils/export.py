"""
题库导出工具
支持多种格式导出：JSON, Markdown, DOCX, PDF, XLSX
"""
import io
import json
from datetime import datetime
from typing import Dict, List, Any

try:
    from docx import Document
    from docx.shared import Inches
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.lib.fonts import addMapping
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, PatternFill
    XLSX_AVAILABLE = True
except ImportError:
    XLSX_AVAILABLE = False


class BankExporter:
    """题库导出器"""
    
    def __init__(self, bank_data: Dict[str, Any]):
        self.bank_data = bank_data
        self.bank_info = bank_data.get('bank_info', {})
        self.questions = bank_data.get('questions', [])

    def _register_chinese_fonts(self):
        """注册中文字体"""
        if not PDF_AVAILABLE:
            return

        try:
            # 尝试注册系统中的中文字体
            import os
            import platform

            # 常见的中文字体路径
            font_paths = []
            system = platform.system()

            if system == "Windows":
                font_paths = [
                    "C:/Windows/Fonts/simsun.ttc",
                    "C:/Windows/Fonts/simhei.ttf",
                    "C:/Windows/Fonts/msyh.ttc"
                ]
            elif system == "Darwin":  # macOS
                font_paths = [
                    "/System/Library/Fonts/PingFang.ttc",
                    "/System/Library/Fonts/STHeiti Light.ttc",
                    "/Library/Fonts/Arial Unicode MS.ttf"
                ]
            else:  # Linux
                font_paths = [
                    "/usr/share/fonts/truetype/wqy/wqy-microhei.ttc",
                    "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc",
                    "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
                    "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
                    "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf"
                ]

            # 尝试注册第一个可用的字体
            for font_path in font_paths:
                if os.path.exists(font_path):
                    try:
                        pdfmetrics.registerFont(TTFont('SimSun', font_path))
                        # 注册字体族
                        addMapping('SimSun', 0, 0, 'SimSun')  # normal
                        addMapping('SimSun', 1, 0, 'SimSun')  # bold
                        addMapping('SimSun', 0, 1, 'SimSun')  # italic
                        addMapping('SimSun', 1, 1, 'SimSun')  # bold-italic
                        return
                    except Exception:
                        continue

            # 如果没有找到中文字体，使用默认字体
            # 这种情况下中文可能显示为方块，但不会报错
            pass

        except Exception:
            # 如果字体注册失败，使用默认字体
            pass
    
    def export_json(self) -> io.BytesIO:
        """导出为JSON格式"""
        json_str = json.dumps(self.bank_data, ensure_ascii=False, indent=2)
        buffer = io.BytesIO(json_str.encode('utf-8'))
        buffer.seek(0)
        return buffer
    
    def export_markdown(self) -> io.BytesIO:
        """导出为Markdown格式"""
        md_content = self._generate_markdown()
        buffer = io.BytesIO(md_content.encode('utf-8'))
        buffer.seek(0)
        return buffer
    
    def export_docx(self) -> io.BytesIO:
        """导出为DOCX格式"""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx 库未安装，无法导出DOCX格式")

        doc = Document()

        # 添加标题
        title = doc.add_heading(self.bank_info.get('name', '题库'), 0)
        title.alignment = 1  # 居中

        # 添加题库信息
        info_table = doc.add_table(rows=0, cols=2)
        info_table.style = 'Table Grid'

        info_items = [
            ('描述', self.bank_info.get('description') or '无'),
            ('分类', self.bank_info.get('category') or '未分类'),
            ('难度', self.bank_info.get('difficulty') or '未设置'),
            ('题目数量', str(self.bank_info.get('question_count', 0))),
            ('导出时间', self.bank_info.get('exported_at') or ''),
            ('导出人', self.bank_info.get('exported_by') or '')
        ]

        for key, value in info_items:
            row = info_table.add_row()
            row.cells[0].text = str(key)
            row.cells[1].text = str(value)

        doc.add_page_break()

        # 添加题目
        questions = self.questions or []  # 确保questions不为None
        for i, question in enumerate(questions, 1):
            if not question:  # 跳过None题目
                continue

            # 题目标题
            doc.add_heading(f'题目 {i}', level=2)

            # 题目信息
            info_p = doc.add_paragraph()
            info_p.add_run(f"类型: {self._get_type_text(question.get('type', 'unknown'))}  ")
            info_p.add_run(f"难度: {question.get('difficulty', '未设置')}  ")
            info_p.add_run(f"分值: {question.get('points', 1)}分")

            # 题目内容
            self._add_question_content_to_docx(doc, question)

            # 答案
            if question.get('answer'):
                doc.add_paragraph('答案:', style='Heading 3')
                answer_text = self._format_answer_text(question)
                doc.add_paragraph(answer_text)

            # 解析
            if question.get('explanation'):
                doc.add_paragraph('解析:', style='Heading 3')
                doc.add_paragraph(question['explanation'])

            doc.add_paragraph()  # 空行

        # 保存到内存
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
    
    def export_pdf(self) -> io.BytesIO:
        """导出为PDF格式"""
        if not PDF_AVAILABLE:
            raise ImportError("reportlab 库未安装，无法导出PDF格式")

        # 注册中文字体
        self._register_chinese_fonts()

        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4)

        # 样式 - 使用中文字体
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'ChineseTitle',
            parent=styles['Heading1'],
            fontName='SimSun',
            fontSize=18,
            spaceAfter=30,
            alignment=1  # 居中
        )

        heading_style = ParagraphStyle(
            'ChineseHeading',
            parent=styles['Heading2'],
            fontName='SimSun',
            fontSize=14,
            spaceAfter=6
        )

        normal_style = ParagraphStyle(
            'ChineseNormal',
            parent=styles['Normal'],
            fontName='SimSun',
            fontSize=10,
            leading=14
        )
        
        story = []
        
        # 标题
        story.append(Paragraph(self.bank_info.get('name', '题库'), title_style))
        story.append(Spacer(1, 12))
        
        # 题库信息表格
        info_data = [
            ['描述', self.bank_info.get('description', '无')],
            ['分类', self.bank_info.get('category', '未分类')],
            ['难度', self.bank_info.get('difficulty', '未设置')],
            ['题目数量', str(self.bank_info.get('question_count', 0))],
            ['导出时间', self.bank_info.get('exported_at', '')],
            ['导出人', self.bank_info.get('exported_by', '')]
        ]
        
        info_table = Table(info_data, colWidths=[2*inch, 4*inch])
        info_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, -1), colors.lightgrey),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, -1), 'SimSun'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
            ('BACKGROUND', (0, 0), (0, -1), colors.grey),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        
        story.append(info_table)
        story.append(Spacer(1, 20))
        
        # 题目
        questions = self.questions or []
        for i, question in enumerate(questions, 1):
            if not question:
                continue

            # 题目标题
            story.append(Paragraph(f'题目 {i}', heading_style))

            # 题目信息
            info_text = f"类型: {self._get_type_text(question.get('type', 'unknown'))} | 难度: {question.get('difficulty', '未设置')} | 分值: {question.get('points', 1)}分"
            story.append(Paragraph(info_text, normal_style))
            story.append(Spacer(1, 6))

            # 题目内容
            content_text = self._format_question_content_for_pdf(question)
            story.append(Paragraph(content_text, normal_style))

            # 答案
            if question.get('answer'):
                story.append(Paragraph('<b>答案:</b>', normal_style))
                answer_text = self._format_answer_text(question)
                story.append(Paragraph(answer_text, normal_style))

            # 解析
            if question.get('explanation'):
                story.append(Paragraph('<b>解析:</b>', normal_style))
                story.append(Paragraph(question['explanation'], normal_style))

            story.append(Spacer(1, 12))
        
        doc.build(story)
        buffer.seek(0)
        return buffer
    
    def export_xlsx(self) -> io.BytesIO:
        """导出为XLSX格式"""
        if not XLSX_AVAILABLE:
            raise ImportError("openpyxl 库未安装，无法导出XLSX格式")
        
        wb = Workbook()
        
        # 题库信息工作表
        info_ws = wb.active
        info_ws.title = "题库信息"
        
        # 设置标题样式
        title_font = Font(size=16, bold=True)
        header_font = Font(bold=True)
        header_fill = PatternFill(start_color="CCCCCC", end_color="CCCCCC", fill_type="solid")
        
        # 题库信息
        info_ws['A1'] = self.bank_info.get('name', '题库')
        info_ws['A1'].font = title_font
        info_ws.merge_cells('A1:B1')
        
        info_data = [
            ('描述', self.bank_info.get('description', '无')),
            ('分类', self.bank_info.get('category', '未分类')),
            ('难度', self.bank_info.get('difficulty', '未设置')),
            ('题目数量', str(self.bank_info.get('question_count', 0))),
            ('导出时间', self.bank_info.get('exported_at', '')),
            ('导出人', self.bank_info.get('exported_by', ''))
        ]
        
        for i, (key, value) in enumerate(info_data, 3):
            info_ws[f'A{i}'] = key
            info_ws[f'B{i}'] = value
            info_ws[f'A{i}'].font = header_font
            info_ws[f'A{i}'].fill = header_fill
        
        # 题目列表工作表
        questions_ws = wb.create_sheet("题目列表")
        
        # 表头
        headers = ['序号', '类型', '题目', '答案', '解析', '难度', '分值']
        for i, header in enumerate(headers, 1):
            cell = questions_ws.cell(row=1, column=i, value=header)
            cell.font = header_font
            cell.fill = header_fill
        
        # 题目数据
        for i, question in enumerate(self.questions, 2):
            questions_ws.cell(row=i, column=1, value=i-1)  # 序号
            questions_ws.cell(row=i, column=2, value=self._get_type_text(question['type']))  # 类型
            questions_ws.cell(row=i, column=3, value=self._format_question_content_for_excel(question))  # 题目
            questions_ws.cell(row=i, column=4, value=self._format_answer_text(question))  # 答案
            questions_ws.cell(row=i, column=5, value=question.get('explanation', ''))  # 解析
            questions_ws.cell(row=i, column=6, value=question.get('difficulty', '未设置'))  # 难度
            questions_ws.cell(row=i, column=7, value=question.get('points', 1))  # 分值
        
        # 调整列宽
        questions_ws.column_dimensions['A'].width = 8
        questions_ws.column_dimensions['B'].width = 12
        questions_ws.column_dimensions['C'].width = 50
        questions_ws.column_dimensions['D'].width = 30
        questions_ws.column_dimensions['E'].width = 40
        questions_ws.column_dimensions['F'].width = 10
        questions_ws.column_dimensions['G'].width = 8
        
        # 保存到内存
        buffer = io.BytesIO()
        wb.save(buffer)
        buffer.seek(0)
        return buffer
    
    def _generate_markdown(self) -> str:
        """生成Markdown格式内容"""
        md_lines = []
        
        # 题库信息
        md_lines.append(f"# {self.bank_info.get('name', '题库')}")
        md_lines.append("")
        
        if self.bank_info.get('description'):
            md_lines.append(f"**描述**: {self.bank_info['description']}")
            md_lines.append("")
        
        md_lines.append(f"**分类**: {self.bank_info.get('category', '未分类')}")
        md_lines.append(f"**难度**: {self.bank_info.get('difficulty', '未设置')}")
        md_lines.append(f"**题目数量**: {self.bank_info.get('question_count', 0)}")
        md_lines.append(f"**导出时间**: {self.bank_info.get('exported_at', '')}")
        md_lines.append(f"**导出人**: {self.bank_info.get('exported_by', '')}")
        md_lines.append("")
        md_lines.append("---")
        md_lines.append("")
        
        # 题目列表
        for i, question in enumerate(self.questions, 1):
            md_lines.append(f"## 题目 {i}")
            md_lines.append("")
            md_lines.append(f"**类型**: {self._get_type_text(question['type'])}")
            md_lines.append(f"**难度**: {question.get('difficulty', '未设置')}")
            md_lines.append(f"**分值**: {question.get('points', 1)} 分")
            md_lines.append("")
            
            # 题目内容
            content_md = self._format_question_content_for_markdown(question)
            md_lines.append(content_md)
            md_lines.append("")
            
            # 答案
            if question.get('answer'):
                md_lines.append("**答案**:")
                answer_text = self._format_answer_text(question)
                md_lines.append(answer_text)
                md_lines.append("")
            
            # 解析
            if question.get('explanation'):
                md_lines.append(f"**解析**: {question['explanation']}")
                md_lines.append("")
            
            md_lines.append("---")
            md_lines.append("")
        
        return "\n".join(md_lines)

    def _get_type_text(self, question_type: str) -> str:
        """获取题目类型文本"""
        type_map = {
            'choice': '选择题',
            'true_false': '判断题',
            'qa': '问答题',
            'math': '数学题',
            'programming': '编程题',
            'fill_blank': '填空题'
        }
        return type_map.get(question_type, question_type)

    def _format_answer_text(self, question: Dict[str, Any]) -> str:
        """格式化答案文本"""
        answer = question.get('answer')
        if not answer:
            return '无'

        question_type = question.get('type')

        if question_type == 'choice':
            if isinstance(answer, dict):
                if answer.get('correct_options'):
                    return f"正确选项: {', '.join(answer['correct_options'])}"
                elif answer.get('correct_option'):
                    return f"正确选项: {answer['correct_option']}"
        elif question_type == 'true_false':
            if isinstance(answer, dict) and 'answer' in answer:
                return '正确' if answer['answer'] else '错误'
            elif isinstance(answer, bool):
                return '正确' if answer else '错误'

        return str(answer)

    def _format_question_content_for_markdown(self, question: Dict[str, Any]) -> str:
        """为Markdown格式化题目内容"""
        content = question.get('content', '')
        question_type = question.get('type')

        if question_type == 'choice':
            if isinstance(content, dict):
                lines = []
                # 处理题干 - 可能在title中或content.question中
                question_text = content.get('question') or question.get('title', '')
                if question_text:
                    lines.append(f"**题干**: {question_text}")
                    lines.append("")

                options = content.get('options')
                if options and isinstance(options, list):
                    lines.append("**选项**:")
                    for option in options:
                        if isinstance(option, dict) and option.get('key') and option.get('text'):
                            lines.append(f"- {option['key']}. {option['text']}")

                return "\n".join(lines) if lines else f"**题干**: {question.get('title', '无题目内容')}"

        elif question_type == 'true_false':
            if isinstance(content, dict) and content.get('question'):
                return f"**题干**: {content['question']}"
            elif isinstance(content, str):
                return f"**题干**: {content}"

        else:
            if isinstance(content, dict) and content.get('question'):
                return f"**题干**: {content['question']}"
            elif isinstance(content, str):
                return f"**题干**: {content}"

        return f"**题干**: {str(content) if content else '无题目内容'}"

    def _format_question_content_for_excel(self, question: Dict[str, Any]) -> str:
        """为Excel格式化题目内容"""
        content = question.get('content', '')
        question_type = question.get('type')

        if question_type == 'choice':
            if isinstance(content, dict):
                parts = []
                # 处理题干 - 可能在title中或content.question中
                question_text = content.get('question') or question.get('title', '')
                if question_text:
                    parts.append(question_text)

                options = content.get('options')
                if options and isinstance(options, list):
                    options_text = []
                    for option in options:
                        if isinstance(option, dict) and option.get('key') and option.get('text'):
                            options_text.append(f"{option['key']}.{option['text']}")
                    if options_text:
                        parts.append(" | ".join(options_text))

                return " | ".join(parts) if parts else question.get('title', '无题目内容')

        elif question_type == 'true_false':
            if isinstance(content, dict) and content.get('question'):
                return content['question']
            elif isinstance(content, str):
                return content

        else:
            if isinstance(content, dict) and content.get('question'):
                return content['question']
            elif isinstance(content, str):
                return content

        return str(content) if content else "无题目内容"

    def _format_question_content_for_pdf(self, question: Dict[str, Any]) -> str:
        """为PDF格式化题目内容"""
        content = question.get('content', '')
        question_type = question.get('type')

        if question_type == 'choice':
            if isinstance(content, dict):
                parts = []
                # 处理题干 - 可能在title中或content.question中
                question_text = content.get('question') or question.get('title', '')
                if question_text:
                    parts.append(f"<b>题干:</b> {question_text}")

                options = content.get('options')
                if options and isinstance(options, list):
                    parts.append("<b>选项:</b>")
                    for option in options:
                        if isinstance(option, dict) and option.get('key') and option.get('text'):
                            parts.append(f"&nbsp;&nbsp;{option['key']}. {option['text']}")

                return "<br/>".join(parts) if parts else f"<b>题干:</b> {question.get('title', '无题目内容')}"

        elif question_type == 'true_false':
            if isinstance(content, dict) and content.get('question'):
                return f"<b>题干:</b> {content['question']}"
            elif isinstance(content, str):
                return f"<b>题干:</b> {content}"

        else:
            if isinstance(content, dict) and content.get('question'):
                return f"<b>题干:</b> {content['question']}"
            elif isinstance(content, str):
                return f"<b>题干:</b> {content}"

        return f"<b>题干:</b> {str(content) if content else '无题目内容'}"

    def _add_question_content_to_docx(self, doc, question: Dict[str, Any]):
        """向DOCX文档添加题目内容"""
        content = question.get('content', '')
        question_type = question.get('type')

        if question_type == 'choice':
            if isinstance(content, dict):
                # 处理题干 - 可能在title中或content.question中
                question_text = content.get('question') or question.get('title', '')
                if question_text:
                    p = doc.add_paragraph()
                    p.add_run('题干: ').bold = True
                    p.add_run(question_text)

                options = content.get('options')
                if options and isinstance(options, list):
                    p = doc.add_paragraph()
                    p.add_run('选项:').bold = True

                    for option in options:
                        if isinstance(option, dict) and option.get('key') and option.get('text'):
                            option_p = doc.add_paragraph(f"{option['key']}. {option['text']}", style='List Bullet')
            else:
                # 如果content不是字典，直接显示为题干
                p = doc.add_paragraph()
                p.add_run('题干: ').bold = True
                p.add_run(str(content) if content else question.get('title', '无题目内容'))

        elif question_type == 'true_false':
            p = doc.add_paragraph()
            p.add_run('题干: ').bold = True
            if isinstance(content, dict) and content.get('question'):
                p.add_run(content['question'])
            elif isinstance(content, str):
                p.add_run(content)
            else:
                p.add_run('无题目内容')

        else:
            p = doc.add_paragraph()
            p.add_run('题干: ').bold = True
            if isinstance(content, dict) and content.get('question'):
                p.add_run(content['question'])
            elif isinstance(content, str):
                p.add_run(content)
            else:
                p.add_run('无题目内容')


def get_available_formats() -> List[str]:
    """获取可用的导出格式"""
    formats = ['json', 'markdown']

    if DOCX_AVAILABLE:
        formats.append('docx')

    if PDF_AVAILABLE:
        formats.append('pdf')

    if XLSX_AVAILABLE:
        formats.append('xlsx')

    return formats
